#!/usr/bin/env python3
"""
Isaac Lab MimicGen Blueprint Technical Report v4 생성 스크립트
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_with_number(doc, text, level):
    """번호가 있는 제목 추가"""
    heading = doc.add_heading(text, level=level)
    return heading


def create_table_with_header(doc, headers, rows):
    """헤더가 있는 테이블 생성"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 헤더
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], "4472C4")
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 데이터
    for row_idx, row_data in enumerate(rows):
        cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            cells[col_idx].text = str(cell_data)

    return table


def main():
    doc = Document()

    # 문서 설정
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)

    # =========================================================================
    # 표지
    # =========================================================================
    title = doc.add_heading('Isaac Lab MimicGen Blueprint', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Technical Report v4', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('합성 로봇 조작 모션 데이터 생성 파이프라인\n\n').bold = True
    info.add_run('NVIDIA Omniverse Isaac Lab + Cosmos AI\n\n')
    info.add_run(f'작성일: {datetime.date.today().strftime("%Y년 %m월 %d일")}\n')
    info.add_run('작성: Claude Code\n')
    info.add_run('버전: v23 (23개 세션 개발 완료)')

    doc.add_page_break()

    # =========================================================================
    # 목차
    # =========================================================================
    doc.add_heading('목차', level=1)

    toc_items = [
        "1. 개요 (Executive Summary)",
        "2. 프로젝트 소개",
        "3. 시스템 아키텍처",
        "4. 개발 환경 구성",
        "5. 핵심 문제와 해결 과정",
        "6. 솔루션 구현 (v1 ~ v23)",
        "7. 개발 세션 상세",
        "8. 실행 가이드",
        "9. 출력 구조",
        "10. 문제 해결 가이드",
        "부록 A: 파일 목록",
        "부록 B: 명령어 참조"
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # =========================================================================
    # 1. Executive Summary
    # =========================================================================
    doc.add_heading('1. 개요 (Executive Summary)', level=1)

    doc.add_heading('1.1 프로젝트 목적', level=2)
    doc.add_paragraph(
        'Isaac Lab MimicGen Blueprint는 NVIDIA Omniverse 기반의 합성 로봇 조작 모션 데이터 생성 '
        '파이프라인입니다. 소수의 인간 시연 데이터를 입력으로 받아 대규모의 다양한 학습 데이터셋을 '
        '생성합니다.'
    )

    doc.add_heading('1.2 핵심 성과', level=2)

    achievements = [
        ('23개 버전 개발', 'v1부터 v23까지 체계적인 버전 관리'),
        ('23개 세션 작업', '상세한 작업 로그와 문제 해결 기록'),
        ('AsyncIO 충돌 해결', 'Jupyter와 Isaac Sim 간 이벤트 루프 충돌 완전 해결'),
        ('Standalone 스크립트', 'Data Generation, Video Encoding 독립 스크립트 구현'),
        ('멀티셀렉트 UI', '다중 카메라/비디오 일괄 처리 지원'),
        ('실행 시간 측정', '파이프라인 단계별 소요 시간 자동 기록'),
    ]

    for title, desc in achievements:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    doc.add_heading('1.3 기술 하이라이트', level=2)

    highlights = [
        'nest_asyncio를 활용한 중첩 이벤트 루프 지원',
        'os.system() 기반 독립 프로세스 실행으로 stdout 리다이렉션 문제 해결',
        'pip_overrides 디렉토리를 통한 패키지 의존성 관리',
        'Warp CUDA 커널을 활용한 GPU 가속 비디오 인코딩',
        'Flask REST API를 통한 Cosmos 서버 비동기 통신',
    ]

    for h in highlights:
        doc.add_paragraph(h, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # 2. 프로젝트 소개
    # =========================================================================
    doc.add_heading('2. 프로젝트 소개', level=1)

    doc.add_heading('2.1 배경', level=2)
    doc.add_paragraph(
        '로봇 조작 학습에는 대량의 시연 데이터가 필요합니다. 하지만 실제 로봇으로 데이터를 '
        '수집하는 것은 시간과 비용이 많이 들고, 안전 문제도 있습니다. 이 프로젝트는 시뮬레이션 '
        '환경에서 합성 데이터를 생성하여 이러한 문제를 해결합니다.'
    )

    doc.add_heading('2.2 3단계 파이프라인', level=2)

    pipeline_table = [
        ('Stage 1', 'Isaac Lab 시뮬레이션', 'Franka 로봇 팔로 큐브 스태킹 시뮬레이션\nSemantic Segmentation, Surface Normals 이미지 생성'),
        ('Stage 2', 'GPU 비디오 인코딩', 'Warp CUDA 커널로 Lambertian Shading 적용\nMP4 비디오 파일 생성'),
        ('Stage 3', 'Cosmos AI 변환', 'NVIDIA Cosmos 모델로 시각적 다양성 증강\n동일한 모션에 다양한 외형 생성'),
    ]

    create_table_with_header(doc, ['단계', '이름', '설명'], pipeline_table)

    doc.add_paragraph()

    doc.add_heading('2.3 하드웨어 요구사항', level=2)

    hw_table = [
        ('로컬 PC', 'Ubuntu 22.04', 'RTX A6000 (48GB VRAM)', 'NVIDIA Driver 535.129.03+'),
        ('Cosmos 노드', 'Ubuntu 22.04', 'H100 (80GB VRAM)', '별도 서버 필요'),
    ]

    create_table_with_header(doc, ['용도', 'OS', 'GPU', '비고'], hw_table)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('중요: ').bold = True
    p.add_run('Isaac Lab과 Cosmos는 동일 머신에서 실행할 수 없습니다. Cosmos는 별도의 H100+ GPU 노드가 필요합니다.')

    doc.add_page_break()

    # =========================================================================
    # 3. 시스템 아키텍처
    # =========================================================================
    doc.add_heading('3. 시스템 아키텍처', level=1)

    doc.add_heading('3.1 전체 구조', level=2)

    arch_text = '''
┌─────────────────────────────────────────────────────────────────┐
│                        로컬 PC (RTX A6000)                        │
│  ┌─────────────────────────────────────────────────────────────┐  │
│  │                   Docker Container                           │  │
│  │  ┌─────────────┐    ┌─────────────┐    ┌─────────────┐     │  │
│  │  │   Jupyter   │───>│ Standalone  │───>│ Isaac Lab   │     │  │
│  │  │   Notebook  │    │  Scripts    │    │ Simulation  │     │  │
│  │  └─────────────┘    └─────────────┘    └─────────────┘     │  │
│  │         │                                    │               │  │
│  │         v                                    v               │  │
│  │  ┌─────────────────────────────────────────────────────┐   │  │
│  │  │              output/ (로컬 마운트)                    │   │  │
│  │  │  ├── {camera}_normals/     ├── videos/              │   │  │
│  │  │  ├── {camera}_segmentation/└── cosmos/              │   │  │
│  │  └─────────────────────────────────────────────────────┘   │  │
│  └─────────────────────────────────────────────────────────────┘  │
└─────────────────────────────────────────────────────────────────┘
                              │
                              │ Flask REST API
                              v
┌─────────────────────────────────────────────────────────────────┐
│                     Cosmos 서버 (H100)                           │
│                   NVIDIA Cosmos AI 모델                          │
└─────────────────────────────────────────────────────────────────┘
'''

    p = doc.add_paragraph()
    p.add_run(arch_text).font.name = 'Courier New'
    p.add_run().font.size = Pt(9)

    doc.add_heading('3.2 주요 파일', level=2)

    files_table = [
        ('notebook/generate_dataset.ipynb', '메인 Jupyter 노트북 - 전체 워크플로우'),
        ('notebook/generate_data_standalone.py', 'Data Generation 독립 실행 스크립트'),
        ('notebook/encode_video_standalone.py', 'Video Encoding 독립 실행 스크립트'),
        ('notebook/notebook_utils.py', 'GPU 가속 비디오 인코딩 함수'),
        ('notebook/notebook_widgets.py', 'Jupyter UI 위젯 (SelectMultiple 등)'),
        ('notebook/app.py', 'Flask API - Cosmos 작업 관리'),
        ('notebook/cosmos_request.py', 'Cosmos 서버 HTTP 클라이언트'),
        ('notebook/stacking_prompt.toml', '씬 설명 프롬프트 템플릿'),
        ('docker-compose.yml', 'Docker 컨테이너 설정'),
        ('launch.sh', '컨테이너 시작 스크립트'),
    ]

    create_table_with_header(doc, ['파일', '설명'], files_table)

    doc.add_page_break()

    # =========================================================================
    # 4. 개발 환경 구성
    # =========================================================================
    doc.add_heading('4. 개발 환경 구성', level=1)

    doc.add_heading('4.1 Docker 설정', level=2)

    doc.add_paragraph('프로젝트는 Docker 컨테이너에서 실행됩니다:')

    docker_code = '''# 컨테이너 시작
xhost +local:
docker compose up -d

# Jupyter 접속
http://localhost:8888/lab/tree/generate_dataset.ipynb

# 컨테이너 종료
docker compose down'''

    p = doc.add_paragraph()
    p.add_run(docker_code).font.name = 'Courier New'

    doc.add_heading('4.2 환경 변수', level=2)

    env_table = [
        ('MODE', 'jupyter / standalone', 'jupyter', '실행 모드'),
        ('NUM_TRIALS', '정수', '1', '데이터 생성 횟수'),
        ('ISAAC_JUPYTER_KERNEL', '1', '1', 'Jupyter 커널 모드'),
        ('VIDEO_CONFIG', 'JSON', '-', '비디오 인코딩 설정'),
    ]

    create_table_with_header(doc, ['변수', '값', '기본값', '설명'], env_table)

    doc.add_heading('4.3 볼륨 마운트', level=2)

    mount_table = [
        ('./datasets', '/workspace/isaaclab/datasets', 'HDF5 데이터셋'),
        ('./output', '/workspace/isaaclab/output', '출력 이미지/비디오'),
        ('./notebook', '/workspace/isaaclab', '노트북 및 스크립트'),
    ]

    create_table_with_header(doc, ['로컬 경로', '컨테이너 경로', '용도'], mount_table)

    doc.add_page_break()

    # =========================================================================
    # 5. 핵심 문제와 해결 과정
    # =========================================================================
    doc.add_heading('5. 핵심 문제와 해결 과정', level=1)

    doc.add_heading('5.1 문제 상황', level=2)
    doc.add_paragraph(
        'Jupyter 노트북에서 Data Generation 셀을 실행하면 "[Trial 1/5] 시작..." 메시지 이후 '
        '무한 대기 상태에 빠졌습니다. 어떤 출력도 없이 [*] 상태로 멈추는 현상이었습니다.'
    )

    doc.add_heading('5.2 근본 원인', level=2)
    doc.add_paragraph(
        'Jupyter와 Isaac Sim(Omniverse Kit)이 각각 자체 asyncio 이벤트 루프를 사용합니다. '
        '두 루프가 충돌하여 Kit의 async_engine에서 다음과 같은 검증이 실패했습니다:'
    )

    error_code = '''assert events._get_running_loop() is self
# Kit의 이벤트 루프와 현재 실행 중인 루프가 다름!'''

    p = doc.add_paragraph()
    p.add_run(error_code).font.name = 'Courier New'

    doc.add_heading('5.3 실패한 접근들', level=2)

    failed_approaches = [
        ('v6', 'loop.run_until_complete()', '블로킹'),
        ('v7', 'ThreadPoolExecutor', '블로킹'),
        ('v8', 'simulation_app.update() 폴링 + 스레드', '블로킹'),
        ('v9', '공식 API (setup_async_generation)', '블로킹 (Jupyter에서)'),
        ('v10', 'subprocess.Popen', '출력이 이전 셀에 표시됨'),
        ('v11', 'subprocess + shell=True', '출력이 이전 셀에 표시됨'),
        ('v12', 'IPython ! 명령', '출력이 이전 셀에 표시됨'),
    ]

    create_table_with_header(doc, ['버전', '접근 방법', '결과'], failed_approaches)

    doc.add_heading('5.4 발견된 핵심 사실', level=2)

    discoveries = [
        'Isaac Sim이 시작되면 해당 셀의 출력 영역에 stdout이 고정됨',
        '이후 셀에서 어떤 방법을 사용해도 출력이 첫 번째 셀에 나타남',
        'subprocess로 별도 프로세스를 실행해도 stdout이 원래 Isaac Sim 셀에 연결됨',
        'IPython 매직 명령(!)은 Isaac Sim 환경에서 지원되지 않음',
    ]

    for d in discoveries:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_heading('5.5 최종 해결책', level=2)
    doc.add_paragraph(
        'Isaac Sim을 Jupyter에서 시작하지 않고, Standalone 스크립트로 별도 프로세스에서 '
        '실행합니다. os.system()을 사용하면 stdout이 현재 셀에 직접 표시됩니다.'
    )

    solution_code = '''# Data Generation 셀 (v17)
SKIP_SIMULATION = True  # Isaac Sim 시작 스킵

# os.system으로 standalone 스크립트 실행
return_code = os.system(
    f"cd /workspace/isaaclab && "
    f"NUM_TRIALS={trials} "
    f"./_isaac_sim/python.sh -u generate_data_standalone.py"
)'''

    p = doc.add_paragraph()
    p.add_run(solution_code).font.name = 'Courier New'

    doc.add_page_break()

    # =========================================================================
    # 6. 솔루션 구현 (v1 ~ v23)
    # =========================================================================
    doc.add_heading('6. 솔루션 구현 (v1 ~ v23)', level=1)

    doc.add_heading('6.1 버전 진화 타임라인', level=2)

    versions = [
        ('v1~v4', '기본 asyncio 접근', '이벤트 루프 충돌'),
        ('v5', 'nest_asyncio + 공식 API', 'Standalone 스크립트에서 성공'),
        ('v6~v8', 'Jupyter에서 직접 실행', '블로킹'),
        ('v9', '공식 API 적용', 'Jupyter에서 블로킹'),
        ('v10~v12', 'subprocess/IPython', 'stdout 리다이렉션 문제'),
        ('v13', 'SKIP_SIMULATION + os.system', '개념 검증'),
        ('v14~v16', 'toml 모듈 문제 해결', 'pip_overrides 사용'),
        ('v17', 'NUM_TRIALS 전달 + os._exit()', '완전 해결'),
        ('v18', '출력 폴더 구조화', '계층적 디렉토리'),
        ('v19', 'Video Preprocessing 호환', '경로 상수 수정'),
        ('v20', 'Video Encoding Standalone', '별도 스크립트'),
        ('v21', '멀티 카메라 UI', 'SelectMultiple 위젯'),
        ('v22', '멀티 비디오 Cosmos', 'SelectMultiple 위젯'),
        ('v23', '실행 시간 측정', 'TIMING_LOG 기능'),
    ]

    create_table_with_header(doc, ['버전', '주요 변경', '결과'], versions)

    doc.add_heading('6.2 핵심 코드 설명', level=2)

    doc.add_heading('6.2.1 generate_data_standalone.py', level=3)

    standalone_code = '''# nest_asyncio 적용 (Jupyter와 동일한 환경)
import nest_asyncio
nest_asyncio.apply()

# 공식 Isaac Lab API 사용
from isaaclab_mimic.datagen.generation import (
    setup_env_config, env_loop, setup_async_generation
)

async_gen = setup_async_generation(
    env=env,
    num_envs=args_cli.num_envs,
    input_file=args_cli.input_file,
    success_term=success_term,
    pause_subtask=args_cli.pause_subtask
)

env_loop(
    env=env,
    env_action_queue=async_gen['action_queue'],
    shared_datagen_info_pool=async_gen['info_pool'],
    asyncio_event_loop=async_gen['event_loop']
)'''

    p = doc.add_paragraph()
    p.add_run(standalone_code).font.name = 'Courier New'

    doc.add_heading('6.2.2 노트북 Data Generation 셀', level=3)

    notebook_code = '''# v17: Standalone 전용 Data Generation
import os, sys

print("[설정] 목표 생성 횟수:", trials)

# 필수 패키지 설치 (pip_overrides에)
os.system(
    "/isaac-sim/kit/python/bin/python3 -m pip install "
    "--target=/workspace/isaaclab/pip_overrides "
    "nest_asyncio toml -q"
)

# standalone 스크립트 실행
return_code = os.system(
    f"cd /workspace/isaaclab && "
    f"NUM_TRIALS={trials} "
    f"./_isaac_sim/python.sh -u generate_data_standalone.py"
)

if return_code == 0:
    print("[SUCCESS] 데이터 생성 완료!")'''

    p = doc.add_paragraph()
    p.add_run(notebook_code).font.name = 'Courier New'

    doc.add_page_break()

    # =========================================================================
    # 7. 개발 세션 상세
    # =========================================================================
    doc.add_heading('7. 개발 세션 상세', level=1)

    sessions = [
        ('세션 1', '2026-02-03', 'Standalone 스크립트 (v5) 성공',
         'asyncio 충돌 분석, nest_asyncio + 공식 API로 해결'),
        ('세션 2', '2026-02-03', 'Git 저장소 설정',
         'origin/upstream 분리, SSH 키 설정'),
        ('세션 3', '2026-02-03', '프로젝트 문서화',
         '3D 시각화 HTML, 기술 문서 DOCX 생성'),
        ('세션 4', '2026-02-03', 'Standalone 모드 개선',
         '볼륨 마운트, NUM_TRIALS 환경 변수'),
        ('세션 5', '2026-02-03', '시각화/문서 v2',
         'project_visualization_v2.html'),
        ('세션 6', '2026-02-04', '노트북 v9 수정',
         '공식 API 적용 시도'),
        ('세션 7', '2026-02-04', 'ISAAC_JUPYTER_KERNEL',
         '환경 변수 추가'),
        ('세션 8', '2026-02-04', 'subprocess 방식 (v10)',
         'Popen으로 standalone 실행 시도'),
        ('세션 9', '2026-02-04', 'subprocess 개선 (v11)',
         'shell=True, start_new_session 시도'),
        ('세션 10', '2026-02-04', 'IPython ! 명령 (v12)',
         '매직 명령 시도'),
        ('세션 11', '2026-02-04', 'SKIP_SIMULATION (v13)',
         'os.system 방식 도입'),
        ('세션 12', '2026-02-04', 'toml 모듈 수정 (v14)',
         '패키지 설치 추가'),
        ('세션 13', '2026-02-04', 'Python 경로 수정 (v15)',
         '올바른 Python 바이너리 사용'),
        ('세션 14', '2026-02-04', 'pip_overrides (v16)',
         '--target 옵션으로 설치 위치 지정'),
        ('세션 15', '2026-02-04', '3가지 개선 (v17)',
         'NUM_TRIALS 직접 전달, os._exit() 사용'),
        ('세션 16', '2026-02-04', '폴더 구조화 (v18)',
         'organize_output_files() 함수'),
        ('세션 17', '2026-02-04', 'Video Preprocessing (v19)',
         '경로 상수 수정, 계층 구조 지원'),
        ('세션 18', '2026-02-04', 'Video Encoding Standalone (v20)',
         'encode_video_standalone.py 생성'),
        ('세션 19', '2026-02-04', '시각화/문서 v3',
         'project_visualization_v3.html, Report v3'),
        ('세션 20', '2026-02-11', '멀티 카메라 UI (v21)',
         'SelectMultiple 위젯으로 변경'),
        ('세션 21', '2026-02-11', '멀티 비디오 Cosmos (v22)',
         'Cosmos 입력도 SelectMultiple'),
        ('세션 22', '2026-02-11', '실행 시간 측정 (v23)',
         'TIMING_LOG 기능 추가'),
        ('세션 23', '2026-02-11', '시각화 v4',
         'project_visualization_v4.html'),
    ]

    create_table_with_header(doc, ['세션', '날짜', '주요 작업', '상세'], sessions)

    doc.add_page_break()

    # =========================================================================
    # 8. 실행 가이드
    # =========================================================================
    doc.add_heading('8. 실행 가이드', level=1)

    doc.add_heading('8.1 사전 요구사항', level=2)

    prereqs = [
        'Docker 및 Docker Compose 설치',
        'NVIDIA Container Toolkit 설치',
        'RTX A6000 또는 동급 GPU (48GB+ VRAM)',
        'Cosmos 서버 접속 가능 (H100 GPU)',
    ]

    for p_item in prereqs:
        doc.add_paragraph(p_item, style='List Bullet')

    doc.add_heading('8.2 실행 순서', level=2)

    steps = [
        ('1. 컨테이너 시작', 'xhost +local: && docker compose up -d'),
        ('2. Jupyter 접속', 'http://localhost:8888'),
        ('3. 노트북 열기', 'generate_dataset.ipynb'),
        ('4. 커널 초기화', 'Kernel → Restart Kernel and Clear Outputs'),
        ('5. 전체 실행', 'Run → Run All Cells'),
    ]

    for step, cmd in steps:
        p = doc.add_paragraph()
        p.add_run(f'{step}: ').bold = True
        p.add_run(cmd).font.name = 'Courier New'

    doc.add_heading('8.3 예상 출력', level=2)

    expected_output = '''============================================================
[NOTEBOOK START] 2026-02-11 14:30:00
============================================================

[v17] Data Generation - Standalone 전용
[설정] 목표 생성 횟수: 3
...
[TIMING] Data Generation: 15.23 minutes

[v21] Video Encoding - Multi-Camera
[1/2] table_cam 인코딩 중...
[2/2] table_high_cam 인코딩 중...
[TIMING] Video Encoding: 8.45 minutes

[v22] Cosmos Processing - Multi-Video
[1/6] table_cam_demo_0.mp4 처리 중...
...
============================================================
[NOTEBOOK COMPLETE] Execution Time Summary
============================================================
  data_generation     :  15.23 minutes
  video_encoding      :   8.45 minutes
  cosmos_processing   :  51.67 minutes
------------------------------------------------------------
  TOTAL               :  75.35 minutes
============================================================'''

    p = doc.add_paragraph()
    p.add_run(expected_output).font.name = 'Courier New'
    p.add_run().font.size = Pt(9)

    doc.add_page_break()

    # =========================================================================
    # 9. 출력 구조
    # =========================================================================
    doc.add_heading('9. 출력 구조', level=1)

    doc.add_heading('9.1 디렉토리 레이아웃', level=2)

    output_structure = '''output/
├── table_cam_normals/           # Surface Normals 이미지
│   ├── demo_0/
│   │   ├── step_0000.png
│   │   ├── step_0001.png
│   │   └── ...
│   ├── demo_1/
│   └── demo_2/
├── table_cam_segmentation/      # Semantic Segmentation 이미지
│   └── ...
├── table_high_cam_normals/      # 두 번째 카메라
│   └── ...
├── table_high_cam_segmentation/
│   └── ...
├── videos/                      # 인코딩된 비디오
│   ├── table_cam_demo_0.mp4
│   ├── table_cam_demo_1.mp4
│   ├── table_high_cam_demo_0.mp4
│   └── ...
└── cosmos/                      # Cosmos 변환 결과
    ├── cosmos_table_cam_demo_0_42.mp4
    └── ...

datasets/
├── annotated_dataset.hdf5       # 입력 (인간 시연)
└── generated_dataset.hdf5       # 출력 (생성된 데이터)'''

    p = doc.add_paragraph()
    p.add_run(output_structure).font.name = 'Courier New'

    doc.add_heading('9.2 파일 명명 규칙', level=2)

    naming_table = [
        ('이미지', '{camera}_{modality}/demo_{N}/step_{NNNN}.png',
         'table_cam_normals/demo_0/step_0001.png'),
        ('비디오', '{camera}_demo_{N}.mp4',
         'table_cam_demo_0.mp4'),
        ('Cosmos', 'cosmos_{camera}_demo_{N}_{seed}.mp4',
         'cosmos_table_cam_demo_0_42.mp4'),
    ]

    create_table_with_header(doc, ['유형', '패턴', '예시'], naming_table)

    doc.add_page_break()

    # =========================================================================
    # 10. 문제 해결 가이드
    # =========================================================================
    doc.add_heading('10. 문제 해결 가이드', level=1)

    troubleshooting = [
        ('ModuleNotFoundError: toml',
         'pip_overrides 디렉토리가 없거나 패키지 미설치',
         'Data Generation 셀의 pip install 명령 확인'),
        ('셀이 [*] 상태로 멈춤',
         'asyncio 이벤트 루프 충돌',
         'SKIP_SIMULATION = True 확인'),
        ('출력이 이전 셀에 표시됨',
         'Isaac Sim stdout 리다이렉션',
         'os.system() 방식 사용'),
        ('FileNotFoundError: output',
         '볼륨 마운트 미설정',
         'docker-compose.yml 확인'),
        ('GPU 메모리 부족',
         '이전 프로세스가 메모리 점유',
         'docker compose down && up -d'),
        ('Cosmos 연결 실패',
         '서버 미실행 또는 네트워크 문제',
         'app.py가 실행 중인지 확인'),
        ('비디오 없음',
         '카메라 선택 안됨',
         'SelectMultiple에서 카메라 선택 확인'),
        ('NUM_TRIALS 적용 안됨',
         '환경 변수 전달 실패',
         'os.system 명령에서 직접 지정'),
    ]

    create_table_with_header(doc, ['문제', '원인', '해결책'], troubleshooting)

    doc.add_page_break()

    # =========================================================================
    # 부록 A: 파일 목록
    # =========================================================================
    doc.add_heading('부록 A: 파일 목록', level=1)

    all_files = [
        ('notebook/generate_dataset.ipynb', '메인 Jupyter 노트북'),
        ('notebook/generate_data_standalone.py', 'Data Generation 스크립트'),
        ('notebook/encode_video_standalone.py', 'Video Encoding 스크립트'),
        ('notebook/notebook_utils.py', '비디오 인코딩 유틸리티'),
        ('notebook/notebook_widgets.py', 'Jupyter 위젯'),
        ('notebook/app.py', 'Cosmos Flask API'),
        ('notebook/cosmos_request.py', 'Cosmos HTTP 클라이언트'),
        ('notebook/stacking_prompt.toml', '프롬프트 템플릿'),
        ('docker-compose.yml', 'Docker 설정'),
        ('launch.sh', '시작 스크립트'),
        ('work_log.txt', '작업 로그'),
        ('CLAUDE.md', 'Claude Code 컨텍스트'),
        ('project_visualization_v4.html', '3D 시각화'),
    ]

    create_table_with_header(doc, ['파일', '설명'], all_files)

    # =========================================================================
    # 부록 B: 명령어 참조
    # =========================================================================
    doc.add_heading('부록 B: 명령어 참조', level=1)

    commands = [
        ('컨테이너 시작', 'xhost +local: && docker compose up -d'),
        ('컨테이너 종료', 'docker compose down'),
        ('컨테이너 재시작', 'docker compose down && docker compose up -d'),
        ('컨테이너 로그', 'docker compose logs -f'),
        ('Standalone 실행', 'MODE=standalone NUM_TRIALS=5 docker compose run --rm isaac-lab'),
        ('비디오 확인', 'ls -la output/videos/'),
        ('비디오 재생', 'vlc output/videos/table_cam_demo_0.mp4'),
    ]

    create_table_with_header(doc, ['작업', '명령어'], commands)

    # =========================================================================
    # 저장
    # =========================================================================
    output_path = '/home/etri/synthetic-manipulation-motion-generation/Isaac_Lab_MimicGen_Technical_Report_v4.docx'
    doc.save(output_path)
    print(f"문서가 저장되었습니다: {output_path}")

    return output_path


if __name__ == "__main__":
    main()
