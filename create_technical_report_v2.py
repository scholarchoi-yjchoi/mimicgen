#!/usr/bin/env python3
"""
Isaac Lab MimicGen Technical Report Generator (v2)
Creates a comprehensive DOCX document for project documentation.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(paragraph):
    """Add a horizontal line below the paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '76B900')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_document():
    doc = Document()

    # Set document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ================================================================
    # TITLE PAGE
    # ================================================================
    title = doc.add_heading('Isaac Lab MimicGen', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Synthetic Manipulation Motion Generation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0] if subtitle.runs else subtitle.add_run()
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(118, 185, 0)

    doc.add_paragraph()

    report_title = doc.add_paragraph('Technical Report')
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_title.runs[0].font.size = Pt(24)
    report_title.runs[0].font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Document info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Project Documentation\n').bold = True
    info.add_run(f'Version 1.0\n')
    info.add_run(f'Date: {datetime.date.today().strftime("%Y-%m-%d")}\n')
    info.add_run('Author: Claude Code (AI Assistant)')

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS
    # ================================================================
    doc.add_heading('Table of Contents', 1)

    toc_items = [
        ('1. Project Overview', '개요'),
        ('2. System Architecture', '시스템 구성'),
        ('3. Development Environment', '개발 환경'),
        ('4. Data Generation Pipeline', '데이터 생성 파이프라인'),
        ('5. Problem Analysis', '문제 분석'),
        ('6. Solution Implementation', '해결책 구현'),
        ('7. Development Sessions', '개발 세션'),
        ('8. Execution Guide', '실행 가이드'),
        ('9. Results and Outputs', '결과물'),
        ('10. Troubleshooting', '문제 해결'),
        ('Appendix A: File Structure', '파일 구조'),
        ('Appendix B: Command Reference', '명령어 참조'),
    ]

    for item, desc in toc_items:
        p = doc.add_paragraph()
        p.add_run(f'{item}').bold = True
        p.add_run(f' - {desc}')

    doc.add_page_break()

    # ================================================================
    # 1. PROJECT OVERVIEW
    # ================================================================
    h1 = doc.add_heading('1. Project Overview', 1)
    add_horizontal_line(h1)

    doc.add_heading('1.1 Introduction', 2)
    doc.add_paragraph(
        'Isaac Lab MimicGen은 NVIDIA Isaac Sim 시뮬레이터 기반의 로봇 매니퓰레이션 '
        '학습 데이터 합성 생성 프로젝트입니다. 소량의 인간 시연(demonstration) 데이터로부터 '
        '대량의 다양한 학습 데이터를 자동으로 생성하여 로봇 학습의 데이터 효율성을 높입니다.'
    )

    doc.add_heading('1.2 Key Components', 2)

    components = [
        ('Isaac Sim', 'NVIDIA의 물리 기반 로봇 시뮬레이터. GPU 가속 물리 엔진과 포토리얼리스틱 렌더링 제공'),
        ('Isaac Lab', 'Isaac Sim 위에서 동작하는 로봇 학습 프레임워크. 강화학습/모방학습 환경 제공'),
        ('MimicGen', '시연 데이터 증폭 시스템. 동작을 분해하고 재조합하여 다양한 변형 생성'),
        ('Franka Emika Panda', '7-DOF 협동 로봇 팔. 정밀한 매니퓰레이션 작업에 적합'),
        ('Docker Container', 'NVIDIA NGC의 공식 Isaac Lab 컨테이너로 재현 가능한 환경 구성'),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Component'
    header_cells[1].text = 'Description'
    set_cell_shading(header_cells[0], '76B900')
    set_cell_shading(header_cells[1], '76B900')
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    for comp, desc in components:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = desc

    doc.add_paragraph()

    doc.add_heading('1.3 Project Goals', 2)
    goals = [
        'Jupyter 노트북의 AsyncIO 블로킹 문제 해결',
        'Standalone 스크립트 방식의 데이터 생성 파이프라인 구축',
        '생성 데이터의 로컬 저장 및 설정 가능한 생성 횟수 지원',
        '재현 가능한 Docker 기반 실행 환경 구성',
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')

    doc.add_page_break()

    # ================================================================
    # 2. SYSTEM ARCHITECTURE
    # ================================================================
    h2 = doc.add_heading('2. System Architecture', 1)
    add_horizontal_line(h2)

    doc.add_heading('2.1 Overall Architecture', 2)
    doc.add_paragraph(
        '시스템은 Docker 컨테이너 내에서 실행되며, 호스트와 볼륨 마운트를 통해 '
        '데이터를 교환합니다. Isaac Sim이 물리 시뮬레이션을 담당하고, '
        'MimicGen이 데이터 증강을 수행합니다.'
    )

    # Architecture diagram (text-based)
    arch_text = '''
┌─────────────────────────────────────────────────────────────┐
│                      Host Machine                            │
│  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐          │
│  │  datasets/  │  │   output/   │  │   scripts/  │          │
│  │  (HDF5)     │  │  (images)   │  │  (Python)   │          │
│  └──────┬──────┘  └──────┬──────┘  └──────┬──────┘          │
│         │                │                │                  │
│         └────────────────┼────────────────┘                  │
│                          │ Docker Volume Mount               │
├──────────────────────────┼──────────────────────────────────┤
│                   Docker Container                           │
│  ┌───────────────────────┴───────────────────────┐          │
│  │              Isaac Sim Runtime                 │          │
│  │  ┌─────────────┐  ┌─────────────────────────┐ │          │
│  │  │  Isaac Lab  │  │      MimicGen           │ │          │
│  │  │  Framework  │◄─┤  Data Augmentation      │ │          │
│  │  └─────────────┘  └─────────────────────────┘ │          │
│  │  ┌─────────────────────────────────────────┐  │          │
│  │  │         Franka Robot Environment        │  │          │
│  │  │         (Cube Stacking Task)            │  │          │
│  │  └─────────────────────────────────────────┘  │          │
│  └───────────────────────────────────────────────┘          │
└─────────────────────────────────────────────────────────────┘
'''

    arch_para = doc.add_paragraph()
    arch_run = arch_para.add_run(arch_text)
    arch_run.font.name = 'Courier New'
    arch_run.font.size = Pt(8)

    doc.add_heading('2.2 Data Flow', 2)

    flow_steps = [
        ('입력', 'annotated_dataset.hdf5 (10개 인간 시연 데이터)'),
        ('처리', 'MimicGen이 동작 분해 및 재조합 수행'),
        ('시뮬레이션', 'Isaac Sim에서 물리 시뮬레이션 실행'),
        ('검증', '성공/실패 판정 (큐브 스태킹 완료 여부)'),
        ('출력', 'generated_dataset.hdf5 + 렌더링 이미지'),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    headers = ['Stage', 'Process', 'Output']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], '76B900')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    for i, (stage, desc) in enumerate(flow_steps):
        row = table.add_row().cells
        row[0].text = f'{i+1}. {stage}'
        row[1].text = desc
        row[2].text = '→' if i < len(flow_steps) - 1 else 'Complete'

    doc.add_page_break()

    # ================================================================
    # 3. DEVELOPMENT ENVIRONMENT
    # ================================================================
    h3 = doc.add_heading('3. Development Environment', 1)
    add_horizontal_line(h3)

    doc.add_heading('3.1 Requirements', 2)

    requirements = [
        ('OS', 'Ubuntu 20.04+ / Windows with WSL2'),
        ('GPU', 'NVIDIA RTX series (CUDA 11.x+)'),
        ('Docker', 'Docker 20.10+ with NVIDIA Container Toolkit'),
        ('Memory', '16GB+ RAM, 8GB+ VRAM'),
        ('Storage', '50GB+ (Docker image ~30GB)'),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Requirement'
    table.rows[0].cells[1].text = 'Specification'
    set_cell_shading(table.rows[0].cells[0], '333333')
    set_cell_shading(table.rows[0].cells[1], '333333')
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    for req, spec in requirements:
        row = table.add_row().cells
        row[0].text = req
        row[1].text = spec

    doc.add_heading('3.2 Docker Configuration', 2)
    doc.add_paragraph('프로젝트는 NVIDIA NGC의 공식 Isaac Lab 컨테이너를 사용합니다:')

    docker_config = '''services:
  isaac-lab:
    image: nvcr.io/nvidia/gr00t-smmg-bp:1.0
    privileged: true
    network_mode: host
    environment:
      ACCEPT_EULA: Y
      MODE: ${MODE:-jupyter}
      NUM_TRIALS: ${NUM_TRIALS:-1}
    runtime: nvidia
    volumes:
      - ./datasets:/workspace/isaaclab/datasets
      - ./output:/workspace/isaaclab/output
      - ./launch.sh:/workspace/isaaclab/launch.sh
      - ./notebook/generate_data_standalone.py:...
    entrypoint: ["/workspace/isaaclab/launch.sh"]'''

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(docker_config)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)

    doc.add_heading('3.3 Environment Variables', 2)

    env_vars = [
        ('MODE', 'jupyter | standalone', '실행 모드 선택'),
        ('NUM_TRIALS', '1, 5, 10, ...', '생성할 성공 데이터 수'),
        ('ACCEPT_EULA', 'Y', 'NVIDIA EULA 동의'),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    for i, header in enumerate(['Variable', 'Values', 'Description']):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], '76B900')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for var, val, desc in env_vars:
        row = table.add_row().cells
        row[0].text = var
        row[1].text = val
        row[2].text = desc

    doc.add_page_break()

    # ================================================================
    # 4. DATA GENERATION PIPELINE
    # ================================================================
    h4 = doc.add_heading('4. Data Generation Pipeline', 1)
    add_horizontal_line(h4)

    doc.add_heading('4.1 Pipeline Overview', 2)
    doc.add_paragraph(
        'MimicGen 파이프라인은 소량의 인간 시연 데이터를 분석하여 '
        '로봇 동작을 서브태스크로 분해하고, 이를 다양한 방식으로 재조합하여 '
        '새로운 학습 데이터를 생성합니다.'
    )

    pipeline_steps = [
        ('1. 데이터 로드', 'annotated_dataset.hdf5에서 10개의 인간 시연 데이터 로드'),
        ('2. 동작 분해', '각 시연을 의미 있는 서브태스크(reach, grasp, lift, place)로 분해'),
        ('3. 변형 생성', '초기 조건과 목표 위치를 변경하여 다양한 시나리오 생성'),
        ('4. 시뮬레이션', 'Isaac Sim에서 생성된 동작 시퀀스 실행'),
        ('5. 검증', '큐브 스태킹 성공 여부 판정'),
        ('6. 저장', '성공한 시연을 HDF5 파일과 렌더링 이미지로 저장'),
    ]

    for step, desc in pipeline_steps:
        p = doc.add_paragraph()
        p.add_run(step + ': ').bold = True
        p.add_run(desc)

    doc.add_heading('4.2 Task: Cube Stacking', 2)
    doc.add_paragraph(
        'Franka 로봇이 테이블 위의 큐브를 집어서 다른 큐브 위에 쌓는 작업입니다. '
        '이 작업은 단순해 보이지만 다음과 같은 도전적인 요소를 포함합니다:'
    )

    challenges = [
        '정밀한 그리퍼 위치 제어',
        '물체 접촉 시 힘 제어',
        '불안정한 스태킹 상태 관리',
        '다양한 큐브 초기 위치 대응',
    ]
    for c in challenges:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_heading('4.3 Output Data Structure', 2)

    doc.add_paragraph('생성되는 HDF5 파일의 구조:')

    hdf5_structure = '''generated_dataset.hdf5
├── data/
│   ├── demo_0/
│   │   ├── actions          (N, 7) - 로봇 동작 명령
│   │   ├── obs/             관측 데이터
│   │   │   ├── joint_pos    (N, 7) - 관절 위치
│   │   │   ├── joint_vel    (N, 7) - 관절 속도
│   │   │   └── ee_pose      (N, 7) - 엔드이펙터 포즈
│   │   ├── states           (N, D) - 시뮬레이션 상태
│   │   └── initial_state    초기 상태
│   ├── demo_1/
│   └── ...
└── env_args                  환경 설정'''

    struct_para = doc.add_paragraph()
    struct_run = struct_para.add_run(hdf5_structure)
    struct_run.font.name = 'Courier New'
    struct_run.font.size = Pt(9)

    doc.add_page_break()

    # ================================================================
    # 5. PROBLEM ANALYSIS
    # ================================================================
    h5 = doc.add_heading('5. Problem Analysis', 1)
    add_horizontal_line(h5)

    doc.add_heading('5.1 Original Issue', 2)
    doc.add_paragraph(
        'Jupyter 노트북에서 Data Generation 셀을 실행하면 "[Trial 1/5] 시작..." '
        '메시지 출력 후 무한 대기 상태에 빠지는 문제가 발생했습니다.'
    )

    doc.add_heading('5.2 Root Cause Analysis', 2)
    doc.add_paragraph('문제의 근본 원인은 AsyncIO 이벤트 루프 충돌입니다:')

    causes = [
        ('Jupyter 이벤트 루프', 'Jupyter는 내부적으로 asyncio 이벤트 루프를 사용하여 셀 실행 관리'),
        ('Isaac Sim (Kit) 이벤트 루프', 'NVIDIA Kit 엔진도 자체 asyncio 이벤트 루프를 사용'),
        ('충돌 지점', 'asyncio.get_event_loop()가 Kit의 루프가 아닌 새 루프를 반환'),
        ('검증 실패', 'Kit의 async_engine에서 "assert events._get_running_loop() is self" 실패'),
    ]

    for cause, desc in causes:
        p = doc.add_paragraph()
        p.add_run(f'{cause}: ').bold = True
        p.add_run(desc)

    doc.add_heading('5.3 Error Message', 2)

    error_msg = '''AssertionError: assert events._get_running_loop() is self

Traceback:
  File "omni/kit/async_engine/async_engine.py", line XX
    assert events._get_running_loop() is self
AssertionError'''

    error_para = doc.add_paragraph()
    error_run = error_para.add_run(error_msg)
    error_run.font.name = 'Courier New'
    error_run.font.size = Pt(9)
    error_run.font.color.rgb = RGBColor(239, 68, 68)

    doc.add_heading('5.4 Failed Attempts', 2)

    failed_attempts = [
        ('v1: 기본 asyncio 루프', 'asyncio.new_event_loop() + run_until_complete()', '__aenter__ 오류'),
        ('v2: 더미 락 클래스', 'DummyAsyncLock으로 async with 지원', 'AssertionError 지속'),
        ('v3: 공식 API 사용', 'setup_async_generation(), env_loop()', '루프 충돌 지속'),
        ('v4: Kit 루프 직접 접근', 'omni.kit.async_engine._loop 접근', 'AttributeError'),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    for i, header in enumerate(['Version', 'Approach', 'Result']):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], 'EF4444')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for ver, approach, result in failed_attempts:
        row = table.add_row().cells
        row[0].text = ver
        row[1].text = approach
        row[2].text = result

    doc.add_page_break()

    # ================================================================
    # 6. SOLUTION IMPLEMENTATION
    # ================================================================
    h6 = doc.add_heading('6. Solution Implementation', 1)
    add_horizontal_line(h6)

    doc.add_heading('6.1 Final Solution: v5', 2)
    doc.add_paragraph(
        '최종 해결책은 nest_asyncio 라이브러리와 Standalone 스크립트 방식의 조합입니다.'
    )

    solution_components = [
        ('nest_asyncio.apply()', '중첩된 이벤트 루프를 허용하여 Jupyter와 Kit이 공존 가능'),
        ('Standalone Script', 'Jupyter 대신 독립 Python 스크립트로 실행하여 환경 충돌 최소화'),
        ('Official API', 'Isaac Lab의 setup_async_generation(), env_loop() 공식 함수 사용'),
        ('Unbuffered Output', 'Python -u 플래그로 실시간 로그 출력'),
    ]

    for comp, desc in solution_components:
        p = doc.add_paragraph()
        p.add_run(f'{comp}: ').bold = True
        p.add_run(desc)

    doc.add_heading('6.2 Core Code', 2)

    core_code = '''# generate_data_standalone.py

# 1. nest_asyncio 적용 (가장 먼저!)
import nest_asyncio
nest_asyncio.apply()

# 2. Isaac Sim AppLauncher 초기화
from isaaclab.app import AppLauncher
app_launcher = AppLauncher(args_cli)
simulation_app = app_launcher.app

# 3. 환경 변수에서 생성 횟수 읽기
num_trials = int(os.environ.get("NUM_TRIALS", "1"))

# 4. 공식 API로 비동기 생성 설정
from isaaclab_mimic.datagen.generation import (
    setup_env_config,
    env_loop,
    setup_async_generation
)

async_gen = setup_async_generation(
    env=env,
    num_envs=args_cli.num_envs,
    input_file=args_cli.input_file,
    success_term=success_term,
    pause_subtask=args_cli.pause_subtask
)

# 5. 데이터 생성 실행
env_loop(
    env=env,
    env_action_queue=async_gen['action_queue'],
    shared_datagen_info_pool=async_gen['info_pool'],
    asyncio_event_loop=async_gen['event_loop']
)'''

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(core_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)

    doc.add_heading('6.3 Why nest_asyncio Works', 2)
    doc.add_paragraph(
        'nest_asyncio는 Python의 asyncio 이벤트 루프를 패치하여 '
        '이미 실행 중인 루프 내에서 또 다른 루프를 실행할 수 있게 합니다. '
        '이를 통해 Jupyter의 루프와 Kit의 루프가 충돌 없이 공존할 수 있습니다.'
    )

    doc.add_page_break()

    # ================================================================
    # 7. DEVELOPMENT SESSIONS
    # ================================================================
    h7 = doc.add_heading('7. Development Sessions', 1)
    add_horizontal_line(h7)

    sessions = [
        {
            'title': 'Session 1: AsyncIO Problem Resolution',
            'date': '2026-02-03',
            'objective': 'Jupyter 노트북의 Data Generation 블로킹 문제 해결',
            'activities': [
                '문제 원인 분석 (Jupyter + Kit 이벤트 루프 충돌)',
                'v1~v4 다양한 접근법 시도 및 실패',
                'v5: nest_asyncio + 공식 API 조합으로 해결',
                'Standalone 스크립트 방식 채택',
            ],
            'files': ['generate_data_standalone.py (신규)', 'docker-compose.yml', 'launch.sh'],
            'result': '데이터 생성 성공 (1회 시도, 1회 성공)',
        },
        {
            'title': 'Session 2: Git Repository Setup',
            'date': '2026-02-03',
            'objective': 'NVIDIA 원본 저장소와 분리된 개인 저장소 설정',
            'activities': [
                'origin을 upstream으로 변경 (NVIDIA 원본)',
                '개인 저장소를 새 origin으로 설정',
                'SSH 키 생성 및 GitHub 등록',
                '첫 커밋 및 push',
            ],
            'files': ['CLAUDE.md (신규)', 'work_log.txt (신규)'],
            'result': 'GitHub 저장소 설정 완료',
        },
        {
            'title': 'Session 3: Documentation & Visualization',
            'date': '2026-02-03',
            'objective': '프로젝트 문서화 및 시각화 자료 생성',
            'activities': [
                'Three.js 기반 3D 애니메이션 HTML 생성',
                '기술 문서 (DOCX) 작성',
                '프로젝트 구조 및 파이프라인 시각화',
            ],
            'files': ['project_visualization.html (신규)', 'Technical_Report.docx (신규)'],
            'result': '시각화 및 문서화 자료 완성',
        },
        {
            'title': 'Session 4: Standalone Mode Improvements',
            'date': '2026-02-03',
            'objective': 'Standalone 모드의 실용성 개선',
            'activities': [
                '출력 데이터 로컬 저장 (볼륨 마운트 추가)',
                'NUM_TRIALS 환경 변수 지원',
                '컨테이너 재시작 방지 (docker compose run --rm)',
                'pip 의존성 자동 설치 (nest_asyncio, toml)',
            ],
            'files': ['docker-compose.yml', 'launch.sh', 'generate_data_standalone.py'],
            'result': 'NUM_TRIALS=3 테스트 성공 (12회 시도, 3회 성공)',
        },
    ]

    for i, session in enumerate(sessions, 1):
        doc.add_heading(f'7.{i} {session["title"]}', 2)

        p = doc.add_paragraph()
        p.add_run('Date: ').bold = True
        p.add_run(session['date'])

        p = doc.add_paragraph()
        p.add_run('Objective: ').bold = True
        p.add_run(session['objective'])

        act_para = doc.add_paragraph()
        act_run = act_para.add_run('Activities:')
        act_run.bold = True
        act_run.italic = True
        for activity in session['activities']:
            doc.add_paragraph(activity, style='List Bullet')

        files_para = doc.add_paragraph()
        files_run = files_para.add_run('Modified/Created Files:')
        files_run.bold = True
        files_run.italic = True
        for f in session['files']:
            doc.add_paragraph(f, style='List Bullet')

        p = doc.add_paragraph()
        p.add_run('Result: ').bold = True
        result_run = p.add_run(session['result'])
        result_run.font.color.rgb = RGBColor(34, 197, 94)

        doc.add_paragraph()

    doc.add_page_break()

    # ================================================================
    # 8. EXECUTION GUIDE
    # ================================================================
    h8 = doc.add_heading('8. Execution Guide', 1)
    add_horizontal_line(h8)

    doc.add_heading('8.1 Prerequisites', 2)

    prereqs = [
        'Docker 및 NVIDIA Container Toolkit 설치',
        'NVIDIA GPU 드라이버 설치 (CUDA 11.x+)',
        'Git clone 완료',
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')

    doc.add_heading('8.2 Initial Setup', 2)

    setup_code = '''# 1. 프로젝트 클론
git clone https://github.com/scholarchoi-yjchoi/mimicgen.git
cd mimicgen

# 2. 필요한 폴더 생성 및 입력 데이터 복사
mkdir -p datasets output
cp samples/annotated_dataset.hdf5 datasets/

# 3. Docker 이미지 pull (최초 1회, ~30GB)
docker compose pull'''

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(setup_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)

    doc.add_heading('8.3 Running Data Generation', 2)

    doc.add_paragraph('기본 실행 (1회 생성):')
    cmd1 = doc.add_paragraph()
    cmd1_run = cmd1.add_run('MODE=standalone docker compose run --rm isaac-lab')
    cmd1_run.font.name = 'Courier New'
    cmd1_run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph('다중 생성 (예: 5회 성공 목표):')
    cmd2 = doc.add_paragraph()
    cmd2_run = cmd2.add_run('MODE=standalone NUM_TRIALS=5 docker compose run --rm isaac-lab')
    cmd2_run.font.name = 'Courier New'
    cmd2_run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph('Jupyter 모드 (기본):')
    cmd3 = doc.add_paragraph()
    cmd3_run = cmd3.add_run('docker compose up')
    cmd3_run.font.name = 'Courier New'
    cmd3_run.font.size = Pt(10)

    doc.add_heading('8.4 Verifying Results', 2)

    verify_code = '''# 생성된 데이터 확인
ls -la datasets/
# expected: annotated_dataset.hdf5, generated_dataset.hdf5

# 렌더링 이미지 확인
ls -la output/
# expected: table_cam_normals/, table_cam_segmentation/, etc.

# HDF5 파일 내용 확인 (Python)
import h5py
with h5py.File('datasets/generated_dataset.hdf5', 'r') as f:
    print(list(f['data'].keys()))  # ['demo_0', 'demo_1', ...]'''

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(verify_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)

    doc.add_page_break()

    # ================================================================
    # 9. RESULTS AND OUTPUTS
    # ================================================================
    h9 = doc.add_heading('9. Results and Outputs', 1)
    add_horizontal_line(h9)

    doc.add_heading('9.1 Test Results', 2)

    results = [
        ('입력 데이터', '10개 인간 시연'),
        ('테스트 명령', 'MODE=standalone NUM_TRIALS=3 docker compose run --rm isaac-lab'),
        ('총 시도 횟수', '12회'),
        ('성공', '3회'),
        ('실패', '9회'),
        ('성공률', '25%'),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Metric'
    table.rows[0].cells[1].text = 'Value'
    set_cell_shading(table.rows[0].cells[0], '22C55E')
    set_cell_shading(table.rows[0].cells[1], '22C55E')
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for metric, value in results:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value

    doc.add_heading('9.2 Output Files', 2)

    output_structure = '''프로젝트 루트/
├── datasets/
│   ├── annotated_dataset.hdf5   (입력, 10개 시연)
│   └── generated_dataset.hdf5   (출력, N개 시연)
│
└── output/
    ├── table_cam_normals/
    │   ├── 0000.png
    │   ├── 0001.png
    │   └── ...
    ├── table_cam_segmentation/
    ├── table_high_cam_normals/
    └── table_high_cam_segmentation/'''

    struct_para = doc.add_paragraph()
    struct_run = struct_para.add_run(output_structure)
    struct_run.font.name = 'Courier New'
    struct_run.font.size = Pt(9)

    doc.add_heading('9.3 Rendered Images', 2)
    doc.add_paragraph(
        '각 성공 시연에 대해 다중 카메라 뷰의 렌더링 이미지가 생성됩니다:'
    )

    image_types = [
        ('Normals', '표면 법선 벡터 시각화 (RGB 인코딩)'),
        ('Segmentation', '객체별 세그멘테이션 마스크'),
    ]

    for img_type, desc in image_types:
        p = doc.add_paragraph()
        p.add_run(f'{img_type}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ================================================================
    # 10. TROUBLESHOOTING
    # ================================================================
    h10 = doc.add_heading('10. Troubleshooting', 1)
    add_horizontal_line(h10)

    issues = [
        {
            'problem': 'Container keeps restarting after completion',
            'cause': 'docker-compose.yml의 restart: unless-stopped 설정',
            'solution': 'docker compose run --rm 사용 (restart 정책 무시, 완료 후 컨테이너 삭제)',
        },
        {
            'problem': 'ModuleNotFoundError: nest_asyncio',
            'cause': 'Isaac Sim Python 환경에 nest_asyncio 미설치',
            'solution': 'launch.sh에서 pip install nest_asyncio 실행',
        },
        {
            'problem': 'Output data not saved to local',
            'cause': 'Docker 볼륨 마운트 미설정',
            'solution': 'docker-compose.yml에 datasets/, output/ 볼륨 마운트 추가',
        },
        {
            'problem': 'pip dependency resolver error',
            'cause': 'Isaac Sim 컨테이너 내 두 개의 Python 환경 충돌',
            'solution': '무시 가능 (시스템 Python과 Isaac Sim Python이 분리되어 있음)',
        },
        {
            'problem': 'GPU not detected',
            'cause': 'NVIDIA Container Toolkit 미설치 또는 설정 오류',
            'solution': 'nvidia-smi 확인, nvidia-container-toolkit 재설치',
        },
    ]

    for i, issue in enumerate(issues, 1):
        doc.add_heading(f'10.{i} {issue["problem"]}', 2)

        p = doc.add_paragraph()
        p.add_run('Cause: ').bold = True
        p.add_run(issue['cause'])

        p = doc.add_paragraph()
        p.add_run('Solution: ').bold = True
        solution_run = p.add_run(issue['solution'])
        solution_run.font.color.rgb = RGBColor(34, 197, 94)

    doc.add_page_break()

    # ================================================================
    # APPENDIX A: FILE STRUCTURE
    # ================================================================
    ha = doc.add_heading('Appendix A: File Structure', 1)
    add_horizontal_line(ha)

    file_structure = '''synthetic-manipulation-motion-generation/
├── docker-compose.yml          # Docker 서비스 정의
├── launch.sh                   # 컨테이너 엔트리포인트
├── CLAUDE.md                   # Claude Code 프로젝트 컨텍스트
├── work_log.txt                # 작업 로그
├── project_visualization.html  # 3D 시각화
│
├── notebook/
│   ├── generate_data_standalone.py  # Standalone 데이터 생성 스크립트
│   ├── generate_dataset.ipynb       # Jupyter 노트북
│   ├── app.py                       # Gradio 앱
│   ├── cosmos_request.py            # Cosmos API 클라이언트
│   ├── notebook_utils.py            # 유틸리티 함수
│   ├── notebook_widgets.py          # UI 위젯
│   └── stacking_prompt.toml         # 프롬프트 설정
│
├── datasets/
│   ├── annotated_dataset.hdf5       # 입력 (인간 시연)
│   └── generated_dataset.hdf5       # 출력 (생성 데이터)
│
├── output/
│   ├── table_cam_normals/           # 렌더링 이미지
│   ├── table_cam_segmentation/
│   └── ...
│
├── samples/
│   └── annotated_dataset.hdf5       # 원본 샘플 데이터
│
└── pip_overrides/                   # 패키지 오버라이드'''

    struct_para = doc.add_paragraph()
    struct_run = struct_para.add_run(file_structure)
    struct_run.font.name = 'Courier New'
    struct_run.font.size = Pt(9)

    # ================================================================
    # APPENDIX B: COMMAND REFERENCE
    # ================================================================
    hb = doc.add_heading('Appendix B: Command Reference', 1)
    add_horizontal_line(hb)

    commands = [
        ('Standalone 모드 실행 (1회)', 'MODE=standalone docker compose run --rm isaac-lab'),
        ('Standalone 모드 실행 (N회)', 'MODE=standalone NUM_TRIALS=N docker compose run --rm isaac-lab'),
        ('Jupyter 모드 실행', 'docker compose up'),
        ('컨테이너 중지', 'docker compose down'),
        ('로그 확인', 'docker compose logs -f'),
        ('컨테이너 쉘 접속', 'docker compose exec isaac-lab bash'),
        ('이미지 갱신', 'docker compose pull'),
        ('Git push', 'git push origin main'),
        ('원본 저장소 업데이트', 'git fetch upstream && git merge upstream/main'),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Purpose'
    table.rows[0].cells[1].text = 'Command'
    set_cell_shading(table.rows[0].cells[0], '333333')
    set_cell_shading(table.rows[0].cells[1], '333333')
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for purpose, cmd in commands:
        row = table.add_row().cells
        row[0].text = purpose
        row[1].text = cmd
        row[1].paragraphs[0].runs[0].font.name = 'Courier New'
        row[1].paragraphs[0].runs[0].font.size = Pt(9)

    # ================================================================
    # SAVE DOCUMENT
    # ================================================================
    output_path = '/home/etri/synthetic-manipulation-motion-generation/Isaac_Lab_MimicGen_Technical_Report_v2.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_document()
