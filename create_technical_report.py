#!/usr/bin/env python3
"""
Isaac Lab MimicGen Project - Technical Report Generator
Creates a comprehensive DOCX technical document for the project
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

def create_technical_report():
    doc = Document()

    # ============================================================
    # Document Styles
    # ============================================================
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 82, 147)

    # Heading 1
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 82, 147)

    # Heading 2
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(46, 116, 181)

    # ============================================================
    # Title Page
    # ============================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('\n\n\n\nIsaac Lab MimicGen\nSynthetic Manipulation Motion Generation')
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 82, 147)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run('\n\nTechnical Report\n프로젝트 기술 문서')
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(f'\n\n\n\n\n버전: 1.0\n작성일: {datetime.now().strftime("%Y-%m-%d")}\n\n')
    info_run.font.size = Pt(12)

    doc.add_page_break()

    # ============================================================
    # Table of Contents (Manual)
    # ============================================================
    doc.add_heading('목차 (Table of Contents)', level=1)

    toc_items = [
        ('1. 프로젝트 개요', 'Project Overview'),
        ('2. 시스템 아키텍처', 'System Architecture'),
        ('3. 개발 환경 설정', 'Development Environment Setup'),
        ('4. 데이터 생성 파이프라인', 'Data Generation Pipeline'),
        ('5. 문제 상황 및 해결', 'Problem Analysis and Solution'),
        ('6. 핵심 코드 설명', 'Core Code Explanation'),
        ('7. 실행 가이드', 'Execution Guide'),
        ('8. 결과 및 출력물', 'Results and Outputs'),
        ('9. 향후 계획', 'Future Work'),
        ('부록 A: 파일 구조', 'Appendix A: File Structure'),
        ('부록 B: 트러블슈팅', 'Appendix B: Troubleshooting'),
    ]

    for item in toc_items:
        p = doc.add_paragraph()
        p.add_run(f'{item[0]} ({item[1]})').font.size = Pt(11)

    doc.add_page_break()

    # ============================================================
    # 1. Project Overview
    # ============================================================
    doc.add_heading('1. 프로젝트 개요 (Project Overview)', level=1)

    doc.add_heading('1.1 프로젝트 소개', level=2)
    doc.add_paragraph(
        'Isaac Lab MimicGen은 NVIDIA의 Isaac Sim 시뮬레이션 플랫폼을 기반으로 '
        '로봇 조작(Manipulation) 데이터를 자동으로 생성하는 파이프라인입니다. '
        '소수의 인간 시연(Human Demonstration)으로부터 대량의 합성 학습 데이터를 '
        '생성하여 로봇 학습의 데이터 효율성을 크게 향상시킵니다.'
    )

    doc.add_heading('1.2 주요 기술 스택', level=2)

    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = 'Table Grid'
    header_cells = tech_table.rows[0].cells
    header_cells[0].text = '기술'
    header_cells[1].text = '버전'
    header_cells[2].text = '용도'

    tech_data = [
        ('NVIDIA Isaac Sim', '4.5', '물리 시뮬레이션 엔진'),
        ('Isaac Lab', 'Latest', '로봇 학습 프레임워크'),
        ('MimicGen', 'Latest', '데이터 자동 생성'),
        ('Docker', 'Latest', '컨테이너 환경'),
        ('Python', '3.10+', '스크립트 실행'),
        ('PyTorch', '2.11.0+', '딥러닝 프레임워크'),
        ('nest_asyncio', 'Latest', '이벤트 루프 관리'),
    ]

    for tech, ver, purpose in tech_data:
        row = tech_table.add_row().cells
        row[0].text = tech
        row[1].text = ver
        row[2].text = purpose

    doc.add_heading('1.3 프로젝트 목표', level=2)
    goals = [
        '인간 시연 데이터로부터 합성 로봇 조작 데이터 자동 생성',
        'Franka 로봇 팔의 큐브 스태킹(Cube Stacking) 태스크 시뮬레이션',
        'HDF5 포맷의 학습 데이터셋 생성',
        '렌더링된 이미지 데이터(normals, segmentation) 생성',
        'Jupyter 노트북과 Standalone 스크립트 두 가지 실행 방식 지원',
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')

    doc.add_page_break()

    # ============================================================
    # 2. System Architecture
    # ============================================================
    doc.add_heading('2. 시스템 아키텍처 (System Architecture)', level=1)

    doc.add_heading('2.1 전체 시스템 구조', level=2)
    doc.add_paragraph(
        '시스템은 Docker 컨테이너 내에서 실행되며, Isaac Sim과 Isaac Lab이 '
        '핵심 시뮬레이션 및 데이터 생성 기능을 담당합니다.'
    )

    arch_text = '''
┌─────────────────────────────────────────────────────────────┐
│                    Docker Container                          │
│  ┌─────────────────────────────────────────────────────┐    │
│  │              NVIDIA Isaac Sim 4.5                    │    │
│  │  ┌─────────────────┐  ┌─────────────────────────┐   │    │
│  │  │   Isaac Lab     │  │     MimicGen            │   │    │
│  │  │  (Environment)  │  │  (Data Generation)      │   │    │
│  │  └─────────────────┘  └─────────────────────────┘   │    │
│  └─────────────────────────────────────────────────────┘    │
│                           │                                  │
│  ┌────────────────────────┴────────────────────────────┐    │
│  │         Python Standalone / Jupyter Notebook         │    │
│  └─────────────────────────────────────────────────────┘    │
└─────────────────────────────────────────────────────────────┘
                            │
                            ▼
            ┌───────────────────────────────┐
            │       Output Data             │
            │  • HDF5 Dataset               │
            │  • PNG Images (normals/seg)   │
            └───────────────────────────────┘
'''

    arch_para = doc.add_paragraph()
    arch_run = arch_para.add_run(arch_text)
    arch_run.font.name = 'Courier New'
    arch_run.font.size = Pt(8)

    doc.add_heading('2.2 데이터 흐름', level=2)

    flow_table = doc.add_table(rows=1, cols=4)
    flow_table.style = 'Table Grid'
    headers = flow_table.rows[0].cells
    headers[0].text = '단계'
    headers[1].text = '입력'
    headers[2].text = '처리'
    headers[3].text = '출력'

    flow_data = [
        ('1. Human Demo', '텔레오퍼레이션', '시연 기록', 'source_demo.hdf5'),
        ('2. Annotation', 'source_demo.hdf5', '서브태스크 주석', 'annotated_dataset.hdf5'),
        ('3. MimicGen', 'annotated_dataset.hdf5', '합성 데이터 생성', 'generated_dataset.hdf5'),
        ('4. Rendering', '시뮬레이션 상태', 'GPU 렌더링', 'PNG 이미지'),
    ]

    for stage, inp, proc, out in flow_data:
        row = flow_table.add_row().cells
        row[0].text = stage
        row[1].text = inp
        row[2].text = proc
        row[3].text = out

    doc.add_page_break()

    # ============================================================
    # 3. Development Environment
    # ============================================================
    doc.add_heading('3. 개발 환경 설정 (Development Environment)', level=1)

    doc.add_heading('3.1 하드웨어 요구사항', level=2)
    hw_table = doc.add_table(rows=1, cols=3)
    hw_table.style = 'Table Grid'
    hw_headers = hw_table.rows[0].cells
    hw_headers[0].text = '구성요소'
    hw_headers[1].text = '최소 사양'
    hw_headers[2].text = '권장 사양'

    hw_data = [
        ('GPU', 'NVIDIA RTX 3080', 'NVIDIA RTX PRO 6000'),
        ('CPU', '8코어', 'AMD Ryzen 9 16코어'),
        ('RAM', '32GB', '128GB'),
        ('Storage', '100GB SSD', '500GB NVMe SSD'),
    ]

    for comp, min_spec, rec_spec in hw_data:
        row = hw_table.add_row().cells
        row[0].text = comp
        row[1].text = min_spec
        row[2].text = rec_spec

    doc.add_heading('3.2 소프트웨어 요구사항', level=2)
    sw_reqs = [
        'Ubuntu 22.04 LTS',
        'Docker Engine 24.0+',
        'NVIDIA Driver 580+',
        'NVIDIA Container Toolkit',
    ]
    for req in sw_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('3.3 프로젝트 설치', level=2)

    install_code = '''
# 1. 저장소 클론
git clone https://github.com/scholarchoi-yjchoi/mimicgen.git
cd mimicgen

# 2. Docker 컨테이너 실행
docker compose up -d

# 3. 컨테이너 접속 확인
docker exec -it synthetic-manipulation-motion-generation-isaac-lab-1 bash
'''

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(install_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)

    doc.add_heading('3.4 Git Remote 구성', level=2)
    doc.add_paragraph(
        '원본 NVIDIA 저장소를 upstream으로, 개인 저장소를 origin으로 설정하여 '
        '원본에 영향을 주지 않고 개발을 진행합니다.'
    )

    remote_table = doc.add_table(rows=1, cols=3)
    remote_table.style = 'Table Grid'
    remote_headers = remote_table.rows[0].cells
    remote_headers[0].text = 'Remote'
    remote_headers[1].text = 'URL'
    remote_headers[2].text = '용도'

    remote_data = [
        ('origin', 'git@github.com:scholarchoi-yjchoi/mimicgen.git', '개인 저장소 (push)'),
        ('upstream', 'github.com/NVIDIA-Omniverse-blueprints/...', '원본 저장소 (참조)'),
    ]

    for name, url, purpose in remote_data:
        row = remote_table.add_row().cells
        row[0].text = name
        row[1].text = url
        row[2].text = purpose

    doc.add_page_break()

    # ============================================================
    # 4. Data Generation Pipeline
    # ============================================================
    doc.add_heading('4. 데이터 생성 파이프라인 (Data Generation Pipeline)', level=1)

    doc.add_heading('4.1 파이프라인 개요', level=2)

    pipeline_steps = [
        ('Human Demonstration Recording',
         '사람이 텔레오퍼레이션으로 로봇 조작 시연을 기록합니다. '
         '이 시연은 원본 데이터(source demo)로 저장됩니다.'),
        ('Subtask Annotation',
         '기록된 시연에 서브태스크 정보(grasp, lift, place 등)를 주석으로 추가합니다. '
         'MimicGen은 이 주석을 기반으로 데이터를 생성합니다.'),
        ('MimicGen Data Generation',
         '주석된 시연을 기반으로 다양한 변형(variation)의 합성 데이터를 자동 생성합니다. '
         '물체 위치, 초기 상태 등을 변경하여 데이터 다양성을 확보합니다.'),
        ('Visual Augmentation (Cosmos)',
         'NVIDIA Cosmos를 사용하여 생성된 데이터에 비주얼 증강을 적용합니다. '
         '(향후 작업 예정)'),
    ]

    for i, (title, desc) in enumerate(pipeline_steps, 1):
        doc.add_heading(f'4.1.{i} {title}', level=2)
        doc.add_paragraph(desc)

    doc.add_heading('4.2 태스크 설명: Cube Stacking', level=2)
    doc.add_paragraph(
        '본 프로젝트에서 사용하는 태스크는 Franka Panda 로봇 팔을 사용한 큐브 스태킹입니다. '
        '로봇은 테이블 위의 큐브들을 집어 올려 순서대로 쌓아야 합니다.'
    )

    task_info = [
        ('태스크 ID', 'Isaac-Stack-Cube-Franka-IK-Rel-Blueprint-Mimic-v0'),
        ('로봇', 'Franka Panda (7 DoF)'),
        ('액션 공간', '6D End-Effector Pose + 1D Gripper'),
        ('관측 공간', 'Joint positions, velocities, object poses, EEF pose'),
        ('서브태스크', 'grasp_1 → stack_1 → grasp_2 → ...'),
    ]

    task_table = doc.add_table(rows=1, cols=2)
    task_table.style = 'Table Grid'
    task_headers = task_table.rows[0].cells
    task_headers[0].text = '항목'
    task_headers[1].text = '값'

    for item, value in task_info:
        row = task_table.add_row().cells
        row[0].text = item
        row[1].text = value

    doc.add_page_break()

    # ============================================================
    # 5. Problem and Solution
    # ============================================================
    doc.add_heading('5. 문제 상황 및 해결 (Problem Analysis and Solution)', level=1)

    doc.add_heading('5.1 문제 상황', level=2)
    doc.add_paragraph(
        'Isaac Lab MimicGen Blueprint 노트북의 "Data Generation" 셀 실행 시 '
        '"[Trial 1/5] 시작..." 메시지 이후 무한 대기 상태에 빠지는 문제가 발생했습니다.'
    )

    doc.add_heading('5.2 근본 원인 분석', level=2)
    doc.add_paragraph(
        'Jupyter 노트북과 Isaac Sim(Kit)이 각각 자체 asyncio 이벤트 루프를 사용하며, '
        '이 두 루프가 충돌하여 블로킹이 발생했습니다.'
    )

    cause_text = '''
문제의 핵심:
┌─────────────────────────────────────┐
│     Jupyter IPykernel Event Loop    │ ← 메인 이벤트 루프 제어
└─────────────────┬───────────────────┘
                  │ (충돌!)
                  ▼
┌─────────────────────────────────────┐
│    Isaac Sim Kit Async Engine       │ ← 시뮬레이션 루프
└─────────────────────────────────────┘

오류 메시지:
AssertionError: events._get_running_loop() is self
'''
    cause_para = doc.add_paragraph()
    cause_run = cause_para.add_run(cause_text)
    cause_run.font.name = 'Courier New'
    cause_run.font.size = Pt(9)

    doc.add_heading('5.3 시도한 해결 방법들', level=2)

    attempts_table = doc.add_table(rows=1, cols=4)
    attempts_table.style = 'Table Grid'
    att_headers = attempts_table.rows[0].cells
    att_headers[0].text = '버전'
    att_headers[1].text = '접근 방식'
    att_headers[2].text = '결과'
    att_headers[3].text = '실패 원인'

    attempts_data = [
        ('v6', 'loop.run_until_complete()', '블로킹', '이벤트 루프 충돌'),
        ('v7', 'ThreadPoolExecutor', '블로킹', '이벤트 루프 충돌'),
        ('v8', 'simulation_app.update() 폴링', '블로킹', '이벤트 루프 충돌'),
        ('v1', 'asyncio.new_event_loop()', 'aenter 오류', 'asyncio_lock=None'),
        ('v2', 'DummyAsyncLock 클래스', 'AssertionError', 'Kit 루프 검증 실패'),
        ('v3', '공식 Isaac Lab API', 'AssertionError', '새 루프와 Kit 충돌'),
        ('v4', 'Kit 루프 직접 접근', 'AttributeError', 'API 불일치'),
        ('v5', 'nest_asyncio + 공식 API', '성공!', '-'),
    ]

    for ver, approach, result, reason in attempts_data:
        row = attempts_table.add_row().cells
        row[0].text = ver
        row[1].text = approach
        row[2].text = result
        row[3].text = reason

    doc.add_heading('5.4 최종 해결책', level=2)
    doc.add_paragraph(
        '최종 해결책은 Standalone Python 스크립트에서 nest_asyncio를 적용하고 '
        'Isaac Lab의 공식 API를 사용하는 것입니다.'
    )

    solution_points = [
        'nest_asyncio.apply(): 중첩된 asyncio 이벤트 루프 허용',
        '공식 API 사용: setup_async_generation(), env_loop()',
        'Standalone 실행: Jupyter 외부에서 독립적으로 실행',
        'Python -u 플래그: unbuffered 출력으로 실시간 로그 확인',
    ]
    for point in solution_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_page_break()

    # ============================================================
    # 6. Core Code Explanation
    # ============================================================
    doc.add_heading('6. 핵심 코드 설명 (Core Code Explanation)', level=1)

    doc.add_heading('6.1 generate_data_standalone.py 구조', level=2)

    code_structure = '''
generate_data_standalone.py
│
├── 1. 경로 설정 (pip_overrides 우선)
│   └── sys.path 및 LD_LIBRARY_PATH 설정
│
├── 2. nest_asyncio 적용
│   └── nest_asyncio.apply()  ← 핵심!
│
├── 3. AppLauncher 초기화
│   └── Isaac Sim 시작
│
├── 4. Isaac Lab 모듈 임포트
│   └── isaaclab_mimic.envs, datagen.generation 등
│
├── 5. 환경 설정
│   └── env_cfg, success_term 설정
│
├── 6. 비동기 생성 설정
│   └── setup_async_generation()  ← 공식 API
│
├── 7. 데이터 생성 루프
│   └── env_loop()  ← 공식 API
│
└── 8. 정리
    └── simulation_app.close()
'''

    struct_para = doc.add_paragraph()
    struct_run = struct_para.add_run(code_structure)
    struct_run.font.name = 'Courier New'
    struct_run.font.size = Pt(9)

    doc.add_heading('6.2 핵심 코드 스니펫', level=2)

    core_code = '''
# nest_asyncio 적용 (노트북과 동일한 환경 구성)
import nest_asyncio
nest_asyncio.apply()

# ... AppLauncher 초기화 ...

# 공식 Isaac Lab API 사용
from isaaclab_mimic.datagen.generation import (
    setup_env_config,
    env_loop,
    setup_async_generation
)

# 비동기 생성 설정
async_gen = setup_async_generation(
    env=env,
    num_envs=args_cli.num_envs,
    input_file=args_cli.input_file,
    success_term=success_term,
    pause_subtask=args_cli.pause_subtask
)

# 데이터 생성 루프 실행
env_loop(
    env=env,
    env_action_queue=async_gen['action_queue'],
    shared_datagen_info_pool=async_gen['info_pool'],
    asyncio_event_loop=async_gen['event_loop']
)
'''

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(core_code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)

    doc.add_heading('6.3 launch.sh 설정', level=2)

    launch_code = '''
#!/bin/bash
MODE=${MODE:-jupyter}  # 기본값: jupyter

if [ "$MODE" = "standalone" ]; then
    echo "=== Standalone Data Generation Mode ==="
    cd /workspace/isaaclab
    ./_isaac_sim/python.sh -u generate_data_standalone.py
else
    echo "=== Jupyter Lab Mode ==="
    # Jupyter Lab 실행...
fi
'''

    launch_para = doc.add_paragraph()
    launch_run = launch_para.add_run(launch_code)
    launch_run.font.name = 'Courier New'
    launch_run.font.size = Pt(9)

    doc.add_page_break()

    # ============================================================
    # 7. Execution Guide
    # ============================================================
    doc.add_heading('7. 실행 가이드 (Execution Guide)', level=1)

    doc.add_heading('7.1 Standalone 모드 실행', level=2)

    standalone_steps = '''
# 방법 1: Docker exec로 직접 실행
docker exec synthetic-manipulation-motion-generation-isaac-lab-1 \\
    /workspace/isaaclab/_isaac_sim/python.sh -u \\
    /workspace/isaaclab/generate_data_standalone.py

# 방법 2: 컨테이너 내부에서 실행
docker exec -it synthetic-manipulation-motion-generation-isaac-lab-1 bash
cd /workspace/isaaclab
./_isaac_sim/python.sh -u generate_data_standalone.py

# 방법 3: Docker Compose 환경 변수 사용
MODE=standalone docker compose up
'''

    standalone_para = doc.add_paragraph()
    standalone_run = standalone_para.add_run(standalone_steps)
    standalone_run.font.name = 'Courier New'
    standalone_run.font.size = Pt(9)

    doc.add_heading('7.2 Jupyter 모드 실행', level=2)

    jupyter_steps = '''
# Docker Compose로 실행
docker compose up -d

# 브라우저에서 접속
http://localhost:8888

# 노트북 파일 열기
generate_dataset.ipynb
'''

    jupyter_para = doc.add_paragraph()
    jupyter_run = jupyter_para.add_run(jupyter_steps)
    jupyter_run.font.name = 'Courier New'
    jupyter_run.font.size = Pt(9)

    doc.add_heading('7.3 설정 파라미터', level=2)

    params_table = doc.add_table(rows=1, cols=4)
    params_table.style = 'Table Grid'
    params_headers = params_table.rows[0].cells
    params_headers[0].text = '파라미터'
    params_headers[1].text = '기본값'
    params_headers[2].text = '설명'
    params_headers[3].text = '위치'

    params_data = [
        ('num_envs', '1', '병렬 환경 수', 'config dict'),
        ('generation_num_trials', '1', '생성할 시도 횟수', 'config dict'),
        ('input_file', 'datasets/annotated_dataset.hdf5', '입력 데모 파일', 'config dict'),
        ('output_file', 'datasets/generated_dataset.hdf5', '출력 데이터셋', 'config dict'),
        ('headless', 'True', '헤드리스 모드', 'args_cli'),
    ]

    for param, default, desc, loc in params_data:
        row = params_table.add_row().cells
        row[0].text = param
        row[1].text = default
        row[2].text = desc
        row[3].text = loc

    doc.add_page_break()

    # ============================================================
    # 8. Results
    # ============================================================
    doc.add_heading('8. 결과 및 출력물 (Results and Outputs)', level=1)

    doc.add_heading('8.1 실행 결과 요약', level=2)

    results_table = doc.add_table(rows=1, cols=2)
    results_table.style = 'Table Grid'
    res_headers = results_table.rows[0].cells
    res_headers[0].text = '항목'
    res_headers[1].text = '값'

    results_data = [
        ('총 시도 횟수', '1회'),
        ('성공', '1'),
        ('실패', '0'),
        ('성공률', '100%'),
        ('HDF5 파일 크기', '186KB'),
        ('생성된 이미지', '912개'),
    ]

    for item, value in results_data:
        row = results_table.add_row().cells
        row[0].text = item
        row[1].text = value

    doc.add_heading('8.2 출력 데이터 구조', level=2)

    output_structure = '''
datasets/
└── generated_dataset.hdf5 (186KB)
    └── demo_0/
        ├── actions          # 로봇 액션 시퀀스
        ├── initial_state    # 초기 상태
        ├── obs              # 관측 데이터
        └── states           # 상태 시퀀스

output/
├── normals/                 # 456개 PNG 이미지
│   ├── table_cam_normals_0000.png
│   ├── table_cam_normals_0001.png
│   └── ...
└── segmentation/            # 456개 PNG 이미지
    ├── table_cam_segmentation_0000.png
    ├── table_cam_segmentation_0001.png
    └── ...
'''

    out_para = doc.add_paragraph()
    out_run = out_para.add_run(output_structure)
    out_run.font.name = 'Courier New'
    out_run.font.size = Pt(9)

    doc.add_heading('8.3 실행 로그 예시', level=2)

    log_example = '''
============================================================
[Standalone] Data Generation 시작 (v5 - nest_asyncio)
============================================================
PyTorch: 2.11.0.dev20260201+cu128
[1/5] AppLauncher 시작 중...
      -> AppLauncher 완료
[2/5] 모듈 임포트 중...
      -> 모듈 임포트 완료
[3/5] 환경 설정 중...
      -> 환경 설정 완료
[4/5] 비동기 생성 설정 중...
Loaded 10 to datagen info pool
      -> 비동기 생성 설정 완료
[5/5] 데이터 생성 시작!
목표: 1개 성공

**************************************************
have 1 successes out of 1 trials so far
have 0 failures out of 1 trials so far
**************************************************
Reached 1 successes/attempts. Exiting.

============================================================
[완료]
      총 시도: 1회
      성공: 1, 실패: 0
============================================================
'''

    log_para = doc.add_paragraph()
    log_run = log_para.add_run(log_example)
    log_run.font.name = 'Courier New'
    log_run.font.size = Pt(8)

    doc.add_page_break()

    # ============================================================
    # 9. Future Work
    # ============================================================
    doc.add_heading('9. 향후 계획 (Future Work)', level=1)

    future_items = [
        ('Cosmos 비주얼 증강',
         'NVIDIA Cosmos를 사용하여 생성된 데이터에 비주얼 증강을 적용하여 '
         '다양한 조명, 텍스처, 배경 조건의 데이터 생성'),
        ('다중 환경 병렬화',
         'num_envs 파라미터를 증가시켜 GPU 활용도를 높이고 '
         '데이터 생성 속도 향상'),
        ('Jupyter 노트북 개선',
         '노트북의 Data Generation 셀도 nest_asyncio + 공식 API 방식으로 '
         '수정하여 두 실행 방식 모두 지원'),
        ('추가 태스크 지원',
         'Cube Stacking 외에 다른 조작 태스크(Pick & Place, Push 등) '
         '데이터 생성 파이프라인 확장'),
    ]

    for title, desc in future_items:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ============================================================
    # Appendix A: File Structure
    # ============================================================
    doc.add_heading('부록 A: 파일 구조 (Appendix A: File Structure)', level=1)

    file_structure = '''
synthetic-manipulation-motion-generation/
│
├── .gitignore                          # Git 제외 파일 설정
├── CLAUDE.md                           # Claude Code 프로젝트 컨텍스트
├── README.md                           # 프로젝트 소개
├── docker-compose.yml                  # Docker Compose 설정
├── launch.sh                           # 실행 스크립트
├── work_log.txt                        # 작업 로그
├── project_visualization.html          # 프로젝트 시각화
│
├── datasets/
│   ├── annotated_dataset.hdf5          # 주석된 입력 데이터
│   └── generated_dataset.hdf5          # 생성된 출력 데이터
│
├── notebook/
│   ├── generate_dataset.ipynb          # Jupyter 노트북
│   └── generate_data_standalone.py     # Standalone 스크립트
│
├── output/
│   ├── normals/                        # Normal map 이미지
│   └── segmentation/                   # Segmentation 이미지
│
└── pip_overrides/                      # Python 패키지 오버라이드 (Git 제외)
'''

    fs_para = doc.add_paragraph()
    fs_run = fs_para.add_run(file_structure)
    fs_run.font.name = 'Courier New'
    fs_run.font.size = Pt(9)

    doc.add_heading('주요 파일 설명', level=2)

    files_table = doc.add_table(rows=1, cols=3)
    files_table.style = 'Table Grid'
    file_headers = files_table.rows[0].cells
    file_headers[0].text = '파일'
    file_headers[1].text = '용도'
    file_headers[2].text = '수정 여부'

    files_data = [
        ('generate_data_standalone.py', 'Standalone 데이터 생성 스크립트', '신규 생성'),
        ('docker-compose.yml', 'Docker 설정 및 볼륨 마운트', '수정'),
        ('launch.sh', 'MODE 환경변수 지원', '수정'),
        ('.gitignore', 'Git 제외 파일 설정', '신규 생성'),
        ('work_log.txt', '작업 로그 기록', '신규 생성'),
    ]

    for fname, purpose, status in files_data:
        row = files_table.add_row().cells
        row[0].text = fname
        row[1].text = purpose
        row[2].text = status

    doc.add_page_break()

    # ============================================================
    # Appendix B: Troubleshooting
    # ============================================================
    doc.add_heading('부록 B: 트러블슈팅 (Appendix B: Troubleshooting)', level=1)

    troubles = [
        ('출력이 보이지 않음',
         'Isaac Sim 환경에서 stdout이 버퍼링될 수 있습니다.',
         'Python 실행 시 -u 플래그를 사용하여 unbuffered 모드로 실행하세요.\n'
         '예: ./_isaac_sim/python.sh -u script.py'),

        ('파일 수정이 반영되지 않음',
         'Docker 볼륨 마운트가 캐시될 수 있습니다.',
         'docker compose down && docker compose up -d로 컨테이너를 재시작하세요.'),

        ('AssertionError: events._get_running_loop() is self',
         'asyncio 이벤트 루프 충돌입니다.',
         'nest_asyncio.apply()를 스크립트 초기에 호출하세요.'),

        ('ModuleNotFoundError: No module named "xyz"',
         'pip_overrides 경로가 설정되지 않았습니다.',
         'sys.path에 /workspace/isaaclab/pip_overrides를 추가하세요.'),

        ('GPU 메모리 부족',
         '시뮬레이션에 GPU 메모리가 부족합니다.',
         'num_envs를 줄이거나, 다른 GPU 프로세스를 종료하세요.'),
    ]

    for i, (problem, cause, solution) in enumerate(troubles, 1):
        doc.add_heading(f'B.{i} {problem}', level=2)
        doc.add_paragraph(f'원인: {cause}')
        doc.add_paragraph(f'해결: {solution}')

    # ============================================================
    # Save Document
    # ============================================================
    output_path = '/home/etri/synthetic-manipulation-motion-generation/Isaac_Lab_MimicGen_Technical_Report.docx'
    doc.save(output_path)
    print(f'Technical report saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_technical_report()
